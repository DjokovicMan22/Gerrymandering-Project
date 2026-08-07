#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Convert the Week 7 markdown paper into a formatted DOCX.")
    parser.add_argument("--name", default="mi_2020")
    return parser.parse_args()


def is_separator(line: str) -> bool:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return bool(cells) and all(c and set(c) <= set(":-") and "-" in c for c in cells)


def parse_table(lines: list[str], start: int):
    if not lines[start].strip().startswith("|"):
        return None, start
    if start + 1 >= len(lines) or not is_separator(lines[start + 1]):
        return None, start
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        if not is_separator(lines[i]):
            rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
        i += 1
    return rows, i


def add_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Menlo"
        else:
            paragraph.add_run(part)


def add_table(doc, rows) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for j, value in enumerate(rows[0]):
        cell = table.rows[0].cells[j]
        cell.text = value
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for j, value in enumerate(row[: len(cells)]):
            cells[j].text = value
    doc.add_paragraph()


def main() -> None:
    args = parse_args()
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as exc:
        raise SystemExit("Install python-docx first: python -m pip install python-docx") from exc

    md_path = Path(f"outputs/reports/week7/{args.name}_final_paper.md")
    if not md_path.exists():
        raise SystemExit(f"Missing {md_path}. Run scripts/build_week7_deliverables.py first.")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(11.5)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buffer: list[str] = []
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph("\n".join(code_buffer))
                for run in p.runs:
                    run.font.name = "Menlo"
                    run.font.size = Pt(8.5)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue
        if not stripped:
            i += 1
            continue

        table, next_i = parse_table(lines, i)
        if table is not None:
            add_table(doc, table)
            i = next_i
            continue

        img = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if img:
            caption, src = img.groups()
            pth = Path(src)
            if pth.exists():
                try:
                    doc.add_picture(str(pth), width=Inches(6.7))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    p = doc.add_paragraph()
                    add_runs(p, f"[Could not embed image: `{pth}`]")
            else:
                p = doc.add_paragraph()
                add_runs(p, f"[Missing image: `{pth}`]")
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cap.runs:
                r.italic = True
                r.font.size = Pt(9)
            i += 1
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, stripped[2:])
            for r in p.runs:
                r.italic = True
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, stripped[2:])
        elif re.match(r"\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s", "", stripped))
        elif stripped.startswith("**Figure."):
            p = doc.add_paragraph(stripped.replace("**", ""))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.italic = True
                r.font.size = Pt(9)
        else:
            p = doc.add_paragraph()
            add_runs(p, stripped)
        i += 1

    out = Path(f"outputs/reports/week7/{args.name}_final_paper.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Saved {out}")


if __name__ == "__main__":
    main()
